from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX


class Output:

  def __init__(self,output_loc,highlight,quiz_name,q_collect):
    self.output_loc=output_loc
    self.highlight=highlight
    self.q_collect=q_collect
    self.quiz_name=quiz_name

  def exec(self):
    docs = Document()
    style = docs.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    mystyle = docs.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)
    docs.add_heading('\t' + '\t' + '\t' + '\t' + self.quiz_name + '\n')

    for index, question in enumerate(self.q_collect):
      run = docs.add_paragraph(style='rtl').add_run(str(index + 1) + '.' + question.title)
      if self.highlight:
          for choice in question.choices:
            if choice.correct:
              choice_run = docs.add_paragraph(style='rtl').add_run(choice.label)
              choice_run.font.bold = True
              choice_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
              docs.add_paragraph(style='rtl').add_run(choice.label)
      else:
        for choice in question.choices:
          docs.add_paragraph(style='rtl').add_run(choice.label)
      run.style = mystyle
      font = run.font
      font.rtl = True
      run.font.underline = True
    docs.save(self.output_loc+'.docx')